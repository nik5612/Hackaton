from docx import Document
from docx.shared import Pt

text_dict = {
    "number_of_contract" : "123/23",
    "title" : "На поставку расходного материала (Сетка-слинг)",
    "company" : 'ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ "КАМА-МЕДИКА"',
    "company2" : 'Государственное бюджетное учреждение здравоохранения Пермского края «Клиническая медико-санитарная часть №1»',
    "dir" : 'директора Агафонова Валерия Алексеевича',
    "dir2" : 'главного врача Михайленко Дениса Валерьевича',
    "IKZ" : '232590410132259060100100010000000244',
    "source" : 'Средства бюджета учреждения',
    "price" : '73 500 (Семьдесят три тысячи пятьсот) рублей, 00 коп.',
    "time" : '10 (десяти) рабочих дней',
    "time2" : '10 (десяти) рабочих дней',
    "date" : '«30» сентября 2023 года',
    "date2" : 'в течение 10 (десяти) рабочих дней',
    "address" : 'г. Пермь, бул. Гагарина, 68 (Аптека).',
    "Название переменной" : "Значение переменной" #шаблон)
}
file_name = "gloria las plagas"

def gen_docx(file_name, text_dict):
  doc = Document()

  style = doc.styles['Normal']
  style.font.name = 'Times New Roman'
  style.font.size = Pt(10)
  secs = doc.sections
  for secs in doc.sections:
    secs.top_margin = Pt(50)
    secs.bottom_margin = Pt(50)
    secs.left_margin = Pt(50)
    secs.right_margin = Pt(50)

  #temp_doc = Document()
  #doc.styles.add_style('List Number', temp_doc.styles['List Number'].type)
  #par2 = doc.add_paragraph()

  par01 = doc.add_paragraph('')
  par01.alignment = 1
  par01.add_run('Договор №' + text_dict["number_of_contract"]).bold = True

  par001 = doc.add_paragraph('')
  par001.alignment = 1
  par001.add_run(text_dict["title"]).bold = True
  par001.add_run(" ")

  #doc.add_page_break() #Разрыв страницы

  par0 = doc.add_paragraph('')
  par0_f = par0.paragraph_format
  par0_f.first_line_indent = Pt(26)
  par0.add_run("г. Пермь")

  par1 = doc.add_paragraph('')
  par1_f = par1.paragraph_format
  par1_f.first_line_indent = Pt(31)
  par1_f.line_spacing = Pt(12)
  par1.add_run(text_dict['company']).bold = True
  par1.add_run(', в лице ')
  par1.add_run(text_dict['dir'])
  par1.add_run(', действующего на основании Устава, именуемый в дальнейшем «Поставщик» и ')
  par1.add_run(text_dict['company2'] + ', ').bold = True
  par1.add_run('в лице ')
  par1.add_run(text_dict['dir2'])
  par1.add_run(', действующего на основании Устава, именуемый в дальнейшем «Заказчик», в соответствии с требованиями п. 4 ч. 1 ст. 93 Федерального закона от 05 апреля 2013 года № 44-ФЗ «О контрактной системе в сфере закупок товаров, работ, услуг для обеспечения государственных и муниципальных нужд», заключили настоящий Договор о нижеследующем:')

  par2 = doc.add_paragraph('')
  par2.alignment = 1
  par2.add_run('1. ПРЕДМЕТ ДОГОВОРА').bold = True

  par3 = doc.add_paragraph('')
  par3_f = par3.paragraph_format
  par3_f.first_line_indent = Pt(31)
  par3_f.line_spacing = Pt(12)
  par3.add_run('1.1. Поставщик обязуется поставить, а Заказчик обязуется оплатить в обусловленный срок – Товар в соответствии с условиями договора и Спецификацией к нему (Приложение №2).')

  par4 = doc.add_paragraph('')
  par4_f = par4.paragraph_format
  par4_f.first_line_indent = Pt(31)
  par4.add_run('1.2. ИКЗ: ')
  par4.add_run(text_dict['IKZ'])

  par5 = doc.add_paragraph('')
  par5_f = par5.paragraph_format
  par5_f.first_line_indent = Pt(31)
  par5.add_run('1.3. Источник финансирования: ')
  par5.add_run(text_dict['source'])

  par6 = doc.add_paragraph('')
  par6.alignment = 1
  par6.add_run('2. ЦЕНА И ПОРЯДОК РАСЧЕТОВ').bold = True

  par7 = doc.add_paragraph('')
  par7_f = par7.paragraph_format
  par7_f.first_line_indent = Pt(31)
  par7.add_run('2.1. Цена настоящего договора ')
  par7.add_run(text_dict['price']).bold = True
  par7.add_run(' без налога (НДС).')

  par8 = doc.add_paragraph('')
  par8_f = par8.paragraph_format
  par8_f.first_line_indent = Pt(31)
  par8_f.line_spacing = Pt(12)
  par8.add_run('2.2. Цена Договора является твердой и определяется на весь срок исполнения Договора, за исключением случаев, установленных в Законе о контрактной системе')

  par9 = doc.add_paragraph('')
  par9_f = par9.paragraph_format
  par9_f.first_line_indent = Pt(31)
  par9_f.line_spacing = Pt(12)
  par9.add_run('2.3. Цена договора включает в себя стоимость Товара, расходы на выполнение Спецификации (Приложение №1) в полном объеме, перевозку.')

  par10 = doc.add_paragraph('')
  par10_f = par10.paragraph_format
  par10_f.first_line_indent = Pt(31)
  par10_f.line_spacing = Pt(12)
  par10.add_run('2.4 Оплата за поставленный товар осуществляется заказчиком безналичным перечислением денежных средств в течение ')
  par10.add_run(text_dict['time']).bold = True
  par10.add_run(' после проведения Заказчиком приемки товара и предоставления Поставщиком надлежащим образом оформленных платежных документов: счета, счет-фактуры и товарной накладной.')

  par11 = doc.add_paragraph('')
  par11_f = par11.paragraph_format
  par11_f.first_line_indent = Pt(31)
  par11.add_run('2.5. Аванс не предусмотрен.')

  par12 = doc.add_paragraph('')
  par12_f = par12.paragraph_format
  par12_f.first_line_indent = Pt(31)
  par12_f.line_spacing = Pt(12)
  par12.add_run('2.6. Сумма по Договору, подлежащая уплате Поставщику, уменьшается на размер налогов, сборов и иных обязательных платежей в бюджеты бюджетной системы Российской Федерации, связанных с оплатой Договора, если в соответствии с законодательством Российской Федерации о налогах и сборах такие налоги, сборы и иные обязательные платежи подлежат уплате в бюджеты бюджетной системы Российской Федерации Заказчиком.	')

  par13 = doc.add_paragraph('')
  par13_f = par13.paragraph_format
  par13_f.first_line_indent = Pt(31)
  par13_f.line_spacing = Pt(12)
  par13.add_run('2.7. Цена Договора может быть изменена, если по предложению Заказчика увеличивается или уменьшается предусмотренное Договором количество Товара не более чем на десять процентов. ')

  par14 = doc.add_paragraph('')
  par14_f = par14.paragraph_format
  par14_f.first_line_indent = Pt(31)
  par14_f.line_spacing = Pt(12)
  par14.add_run('При этом по соглашению Сторон допускается изменение с учетом положений бюджетного законодательства Российской Федерации цены Договора пропорционально дополнительному количеству Товара исходя из установленной в Договоре цены единицы Товара, но не более чем на десять процентов цены Договора. При уменьшении предусмотренного Договором количества Товара Стороны Договора обязаны уменьшить цену Договора исходя из цены единицы Товара.')

  par15 = doc.add_paragraph('')
  par15_f = par15.paragraph_format
  par15_f.first_line_indent = Pt(31)
  par15_f.line_spacing = Pt(12)
  par15.add_run('Цена единицы дополнительно поставляемого Товара или цена единицы Товара при уменьшении предусмотренного Договором количества поставляемого Товара должна определяться как частное от деления первоначальной цены Договора на предусмотренное в Договоре количество Товара.')

  par16 = doc.add_paragraph('')
  par16.alignment = 1
  par16.add_run('3. СРОКИ ДЕЙСТВИЯ ДОГОВОРА').bold = True

  par17 = doc.add_paragraph('')
  par17_f = par17.paragraph_format
  par17_f.first_line_indent = Pt(31)
  par17_f.line_spacing = Pt(12)
  par17.add_run('3.1. Договор вступает в силу со дня его подписания обеими Сторонами. Срок действия договора c момента заключения Договора до ')
  par17.add_run(text_dict['date']).bold = True
  par17.add_run(', в части расчетов до полного исполнения своих обязательств.')

  par18 = doc.add_paragraph('')
  par18.alignment = 1
  par18.add_run('4. СРОКИ И ПОРЯДОК ПОСТАВКИ ТОВАРА').bold = True

  par19 = doc.add_paragraph('')
  par19_f = par19.paragraph_format
  par19_f.first_line_indent = Pt(31)
  par19.add_run('4.1. Поставка Товара осуществляется со склада Поставщика транспортом Поставщика.')
  
  par20 = doc.add_paragraph('')
  par20_f = par20.paragraph_format
  par20_f.first_line_indent = Pt(31)
  par20_f.line_spacing = Pt(12)
  par20.add_run('4.2. Поставщик осуществляет передачу в течение ')
  par20.add_run(text_dict['date2']).bold = True
  par20.add_run(' со дня заключения контракта. Поставщик за 2 (два) дня до момента поставки товара информирует Заказчика о предстоящей поставке.')

  par21 = doc.add_paragraph('')
  par21_f = par21.paragraph_format
  par21_f.first_line_indent = Pt(31)
  par21.add_run('4.3. Проверка качества Товара производится Заказчиком при его получении от Поставщика.')

  par22 = doc.add_paragraph('')
  par22_f = par22.paragraph_format
  par22_f.first_line_indent = Pt(31)
  par22_f.line_spacing = Pt(12)
  par22.add_run('4.4. Доставка Товара осуществляется силами Поставщика на склад Заказчика по адресу: ')
  par22.add_run(text_dict['address']).bold = True

  
  
  #################################################################
  #  par1 = doc.add_paragraph('')
  #  par2 = doc.add_paragraph('Это третий абзац.')
  # # # добавляем текст во второй параграф
  #  par1.add_run(' Этот текст был добавлен во второй абзац.')
  #  # добавляем текст в третий параграф
  #  par2.add_run(' Добавляем текст в третий абзац.').bold = True
  ###################################################################

  # # добавляем таблицу 3x3
  # table = doc.add_table(rows = 3, cols = 3)
  # # применяем стиль для таблицы
  # table.style = 'Table Grid'

  # # заполняем таблицу данными
  # for row in range(3):
  #     for col in range(3):
  #         # получаем ячейку таблицы
  #         cell = table.cell(row, col)
  #         # записываем в ячейку данные
  #         cell.text = str(row + 1) + str(col + 1)

  for paragraph in doc.paragraphs:
    paragraph.style = 'Normal'
  doc.save('./contracts/' + file_name + '.docx')
